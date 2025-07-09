"""
Утилиты для загрузки и обработки документов Word с расширенной информацией
"""
from docx import Document
from docx.shared import Pt, Cm
from typing import List, Dict, Optional
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import Counter
from docx.oxml import parse_xml
from docx.oxml.ns import qn

class DocumentLoader:
    """Класс для загрузки документов Word с сохранением форматирования"""

    @staticmethod
    def _get_alignment_name(alignment) -> str:
        """Получение названия выравнивания"""
        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: 'LEFT',
            WD_ALIGN_PARAGRAPH.CENTER: 'CENTER',
            WD_ALIGN_PARAGRAPH.RIGHT: 'RIGHT',
            WD_ALIGN_PARAGRAPH.JUSTIFY: 'JUSTIFY',
            WD_ALIGN_PARAGRAPH.DISTRIBUTE: 'DISTRIBUTE'
        }
        return alignment_map.get(alignment, 'LEFT')

    @staticmethod
    def _get_alignment_from_xml(element):
        """Получение выравнивания из XML элемента"""
        try:
            # Проверяем свойства параграфа
            ppr = element.find('.//w:pPr', element.nsmap)
            if ppr is not None:
                jc = ppr.find('.//w:jc', ppr.nsmap)
                if jc is not None and 'w:val' in jc.attrib:
                    align_val = jc.attrib['w:val']
                    alignment_map = {
                        'left': WD_ALIGN_PARAGRAPH.LEFT,
                        'center': WD_ALIGN_PARAGRAPH.CENTER,
                        'right': WD_ALIGN_PARAGRAPH.RIGHT,
                        'both': WD_ALIGN_PARAGRAPH.JUSTIFY,
                        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
                        'distribute': WD_ALIGN_PARAGRAPH.DISTRIBUTE
                    }
                    return alignment_map.get(align_val, None)
        except:
            pass
        return None

    @staticmethod
    def _get_alignment_from_style(style):
        """Получение выравнивания из стиля"""
        try:
            if hasattr(style, 'paragraph_format') and style.paragraph_format:
                if hasattr(style.paragraph_format, 'alignment') and style.paragraph_format.alignment is not None:
                    return style.paragraph_format.alignment
        except:
            pass
        return None

    @staticmethod
    def load_document_with_formatting(file_path: str) -> Dict:
        """Загрузка документа с полной информацией о форматировании"""
        try:
            doc = Document(file_path)

            # Получаем тему документа для определения шрифтов по умолчанию
            theme_fonts = DocumentLoader._get_theme_fonts(doc)

            # Убираем принудительную замену Calibri
            default_font = theme_fonts.get('minor', {}).get('latin', 'Times New Roman')

            document_info = {
                'paragraphs': [],
                'document_properties': DocumentLoader._get_document_properties(doc),
                'page_count': DocumentLoader._estimate_page_count(doc),
                'default_font': default_font,
                'theme_fonts': theme_fonts,
                'styles_info': DocumentLoader._get_styles_info(doc)
            }

            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    para_info = DocumentLoader._extract_paragraph_info(para, i, document_info)

                    document_info['paragraphs'].append(para_info)

            return document_info
        except Exception as e:
            print(f"Ошибка чтения файла: {e}")
            return {'paragraphs': [], 'document_properties': {}, 'page_count': 0, 'default_font': 'Times New Roman'}

    @staticmethod
    def _get_theme_fonts(doc) -> Dict:
        """Получение шрифтов из темы документа, но по умолчанию Times New Roman"""
        theme_fonts = {
            'major': {'latin': 'Times New Roman'},
            'minor': {'latin': 'Times New Roman'},
        }

        try:
            # Получаем XML темы документа
            theme_part = None
            for rel in doc.part.rels.values():
                if rel.reltype.endswith('theme'):
                    theme_part = rel.target_part
                    break

            if theme_part:
                from docx.oxml import parse_xml
                theme_xml = parse_xml(theme_part.blob)

                # Шрифты для заголовков (major)
                major_font = theme_xml.find('.//a:majorFont/a:latin', theme_xml.nsmap)
                if major_font is not None and 'typeface' in major_font.attrib:
                    theme_fonts['major']['latin'] = major_font.attrib['typeface']

                # Шрифты для основного текста (minor)
                minor_font = theme_xml.find('.//a:minorFont/a:latin', theme_xml.nsmap)
                if minor_font is not None and 'typeface' in minor_font.attrib:
                    theme_fonts['minor']['latin'] = minor_font.attrib['typeface']
        except Exception as e:
            print(f"Ошибка при получении шрифтов темы: {e}")

        # Принудительно устанавливаем TNR по умолчанию
        for key in ['major', 'minor']:
            font_val = theme_fonts.get(key, {}).get('latin', '').lower()
            if 'calibri' in font_val:
                theme_fonts[key]['latin'] = 'Times New Roman'

        print("⚠️ theme_fonts = ", theme_fonts)
        return theme_fonts

    @staticmethod
    def _get_styles_info(doc) -> Dict:
        """Получение информации о стилях документа с извлечением шрифтов из XML"""
        styles_info = {}
        try:
            for style in doc.styles:
                if hasattr(style, 'font') and style.font:
                    font_info = {
                        'name': getattr(style.font, 'name', None),
                        'size': getattr(style.font, 'size', None),
                        'bold': getattr(style.font, 'bold', None),
                        'italic': getattr(style.font, 'italic', None)
                    }

                    # Распознаем имя шрифта из XML, если font.name не задан
                    if not font_info['name']:
                        rfonts = style.element.find('.//w:rFonts', style.element.nsmap)
                        if rfonts is not None:
                            for attr in ['w:ascii', 'w:hAnsi', 'w:cs']:
                                if attr in rfonts.attrib:
                                    font_info['name'] = rfonts.attrib[attr]
                                    break

                    if font_info['size']:
                        font_info['size'] = font_info['size'].pt
                    styles_info[style.name] = font_info
        except Exception as e:
            print(f"Ошибка при получении информации о стилях: {e}")

        return styles_info

    @staticmethod
    def _resolve_font_name(raw_font: str, theme_fonts: Dict, default_font: str) -> str:
        """Разрешает шрифт типа minorHAnsi через тему"""
        theme_map = {
            'minorHAnsi': 'minor',
            'majorHAnsi': 'major',
            'minorAscii': 'minor',
            'majorAscii': 'major'
        }
        theme_key = theme_map.get(raw_font)
        if theme_key:
            return theme_fonts.get(theme_key, {}).get('latin', default_font)
        return raw_font


    @staticmethod
    def _get_font_from_xml(element, theme_fonts: Dict, default_font: str) -> Optional[str]:
        """Получение шрифта напрямую из XML элемента с учетом raw и resolved"""
        try:
            rpr = element.find('.//w:rPr', element.nsmap)
            if rpr is not None:
                rfonts = rpr.find('.//w:rFonts', rpr.nsmap)
                if rfonts is not None:
                    for attr in ['w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia']:
                        if attr in rfonts.attrib:
                            raw_font = rfonts.attrib[attr]
                            return raw_font  # raw, потом резолвится отдельно
        except:
            pass
        return None

    @staticmethod
    def _extract_paragraph_info(para, index: int, document_info: Dict) -> Dict:
        """Извлечение информации о параграфе с точным определением шрифта"""
        default_font = document_info['default_font']
        theme_fonts = document_info['theme_fonts']
        styles_info = document_info['styles_info']

        # Собираем информацию о всех runs в параграфе
        runs_data = []
        for run in para.runs:
            if not run.text.strip():
                continue

            run_data = {
                'text': run.text,
                'font': None,
                'size': None,
                'bold': None,
                'italic': None,
                'style': None
            }

            # 1. Проверяем прямое форматирование в run
            if run.font:
                if run.font.name:
                    run_data['font'] = run.font.name
                if run.font.size:
                    run_data['size'] = run.font.size.pt
                run_data['bold'] = run.font.bold
                run_data['italic'] = run.font.italic

             # 2. Если шрифт не найден, проверяем XML элемент run
            if not run_data['font']:
                raw_font = DocumentLoader._get_font_from_xml(run._element, theme_fonts, default_font)
                if raw_font:
                    run_data['font_raw'] = raw_font
                    run_data['font_resolved'] = DocumentLoader._resolve_font_name(raw_font, theme_fonts,
                                                                                  default_font)
                    run_data['font'] = run_data['font_resolved']

            # 3. Проверяем стиль run
            if not run_data['font'] and hasattr(run, 'style') and run.style:
                style_name = run.style.name
                if style_name in styles_info:
                    style_font_info = styles_info[style_name]
                    if style_font_info['name']:
                        run_data['font_raw'] = f"style:{style_name}"
                        run_data['font_resolved'] = style_font_info['name']
                        run_data['font'] = style_font_info['name']
                        run_data['style'] = style_name
                    if style_font_info['size'] and not run_data['size']:
                        run_data['size'] = style_font_info['size']
                    if run_data['bold'] is None:
                        run_data['bold'] = style_font_info['bold']
                    if run_data['italic'] is None:
                        run_data['italic'] = style_font_info['italic']

            # 4.1 Вставляем default Times New Roman, если ничего не найдено
            if not run_data['font']:
                run_data['font_raw'] = 'default'
                run_data['font_resolved'] = default_font
                run_data['font'] = default_font

            if run_data['font']:
                run_data['font_resolved'] = DocumentLoader._resolve_font_name(run_data['font'], theme_fonts,
                                                                              default_font)
                run_data['font_raw'] = run_data['font']
                run_data['font'] = run_data['font_resolved']
            runs_data.append(run_data)

        # 4.2 Проверяем стиль параграфа, если runs не дали результата
        para_style_font = None
        para_style_size = None
        para_style_bold = None
        para_style_italic = None

        if para.style:
            style_name = para.style.name
            if style_name in styles_info:
                style_font_info = styles_info[style_name]
                para_style_font = style_font_info['name']
                para_style_size = style_font_info['size']
                para_style_bold = style_font_info['bold']
                para_style_italic = style_font_info['italic']

        # 5. Проверяем XML параграфа для получения шрифта
        if not para_style_font:
            xml_font = DocumentLoader._get_font_from_xml(para._element, theme_fonts, default_font)
            if xml_font:
                para_style_font = xml_font



        # Заполняем пропущенные значения для runs
        for run_data in runs_data:
            if not run_data['font']:
                if para_style_font:
                    run_data['font'] = para_style_font
                elif para.style and 'Heading' in para.style.name:
                    run_data['font'] = theme_fonts['major'].get('latin', default_font)
                else:
                    run_data['font'] = theme_fonts['minor'].get('latin', default_font)

            if not run_data['size']:
                if para_style_size:
                    run_data['size'] = para_style_size
                else:
                    run_data['size'] = 12.0

            if run_data['bold'] is None:
                run_data['bold'] = para_style_bold if para_style_bold is not None else False

            if run_data['italic'] is None:
                run_data['italic'] = para_style_italic if para_style_italic is not None else False

        # Определяем основные характеристики параграфа
        font_counter = Counter()
        size_counter = Counter()
        bold_chars = 0
        italic_chars = 0
        total_chars = 0

        for run in runs_data:
            text_len = len(run['text'])
            if run['font']:
                font_counter[run['font']] += text_len
            if run['size']:
                size_counter[run['size']] += text_len
            if run['bold']:
                bold_chars += text_len
            if run['italic']:
                italic_chars += text_len
            total_chars += text_len

        # Основной шрифт параграфа
        if font_counter:
            font_name = font_counter.most_common(1)[0][0]
        elif para_style_font:
            font_name = para_style_font
        else:
            font_name = default_font

        # Основной размер шрифта
        if size_counter:
            font_size = size_counter.most_common(1)[0][0]
        elif para_style_size:
            font_size = para_style_size
        else:
            font_size = 12.0

        # Жирность и курсив
        is_bold = bold_chars > total_chars / 2 if total_chars > 0 else False
        is_italic = italic_chars > total_chars / 2 if total_chars > 0 else False

        # Обработка отступов и выравнивания
        left_indent = None
        first_line_indent = None

        try:
            if para.paragraph_format.left_indent:
                left_indent = round(para.paragraph_format.left_indent.cm, 2)
        except:
            pass

        try:
            if para.paragraph_format.first_line_indent:
                first_line_indent = round(para.paragraph_format.first_line_indent.cm, 2)
        except:
            pass

        # Выравнивание - используем несколько методов для надежности
        alignment = WD_ALIGN_PARAGRAPH.LEFT
        alignment_name = 'LEFT'

        try:
            # Метод 1: Проверяем paragraph_format.alignment
            if para.paragraph_format.alignment is not None:
                alignment = para.paragraph_format.alignment
                alignment_name = DocumentLoader._get_alignment_name(alignment)
        except:
            pass

        # Метод 2: Проверяем XML параграфа
        if alignment == WD_ALIGN_PARAGRAPH.LEFT:
            try:
                xml_alignment = DocumentLoader._get_alignment_from_xml(para._element)
                if xml_alignment:
                    alignment = xml_alignment
                    alignment_name = DocumentLoader._get_alignment_name(alignment)
            except:
                pass

        # Метод 3: Проверяем стиль параграфа
        if alignment == WD_ALIGN_PARAGRAPH.LEFT and para.style:
            try:
                style_alignment = DocumentLoader._get_alignment_from_style(para.style)
                if style_alignment:
                    alignment = style_alignment
                    alignment_name = DocumentLoader._get_alignment_name(alignment)
            except:
                pass
        #временно дл проверки:
        # print(f"📌 PARA {index} STYLE = {para.style.name}")
        # print(f"📌 STYLE INFO = {styles_info.get(para.style.name)}")
        return {
            'index': index,
            'text': para.text.strip(),
            'alignment': alignment,
            'alignment_name': alignment_name,
            'font_name': font_name,
            'font_size': round(font_size, 1),
            'is_bold': is_bold,
            'is_italic': is_italic,
            'left_indent': left_indent,
            'first_line_indent': first_line_indent,
            'runs_info': runs_data,
            'style_name': para.style.name if para.style else None,
            'debug': {
                'total_chars': total_chars,
                'unique_fonts': list(font_counter.keys()),
                'font_distribution': dict(font_counter.most_common()),
                'size_distribution': dict(size_counter.most_common()),
                'para_style_font': para_style_font,
                'xml_analysis': DocumentLoader._analyze_paragraph_xml(para),
                'alignment_debug': {
                    'paragraph_format': para.paragraph_format.alignment if hasattr(para.paragraph_format, 'alignment') else None,
                    'xml_alignment': DocumentLoader._get_alignment_from_xml(para._element),
                    'style_alignment': DocumentLoader._get_alignment_from_style(para.style) if para.style else None
                }
            }
        }

    @staticmethod
    def _analyze_paragraph_xml(para) -> Dict:
        """Анализ XML параграфа для отладки"""
        analysis = {
            'has_pPr': False,
            'has_rPr': False,
            'fonts_found': []
        }

        try:
            # Проверяем свойства параграфа
            ppr = para._element.find('.//w:pPr', para._element.nsmap)
            if ppr is not None:
                analysis['has_pPr'] = True
                rpr = ppr.find('.//w:rPr', ppr.nsmap)
                if rpr is not None:
                    analysis['has_rPr'] = True
                    rfonts = rpr.find('.//w:rFonts', rpr.nsmap)
                    if rfonts is not None:
                        for attr in rfonts.attrib:
                            analysis['fonts_found'].append(f"{attr}: {rfonts.attrib[attr]}")

            # Проверяем runs
            for run in para.runs:
                rpr = run._element.find('.//w:rPr', run._element.nsmap)
                if rpr is not None:
                    rfonts = rpr.find('.//w:rFonts', rpr.nsmap)
                    if rfonts is not None:
                        for attr in rfonts.attrib:
                            font_info = f"run_{attr}: {rfonts.attrib[attr]}"
                            if font_info not in analysis['fonts_found']:
                                analysis['fonts_found'].append(font_info)
        except Exception as e:
            analysis['error'] = str(e)

        #временно для провекри
        #print("📄 XML fonts in para:", analysis['fonts_found'])

        return analysis

    @staticmethod
    def _get_document_properties(doc) -> Dict:
        """Получение свойств документа"""
        properties = {}
        try:
            sections = doc.sections
            if sections:
                section = sections[0]
                properties.update({
                    'top_margin': round(section.top_margin.cm, 2) if section.top_margin else None,
                    'bottom_margin': round(section.bottom_margin.cm, 2) if section.bottom_margin else None,
                    'left_margin': round(section.left_margin.cm, 2) if section.left_margin else None,
                    'right_margin': round(section.right_margin.cm, 2) if section.right_margin else None,
                    'page_width': round(section.page_width.cm, 2) if section.page_width else None,
                    'page_height': round(section.page_height.cm, 2) if section.page_height else None
                })
        except Exception as e:
            print(f"Ошибка при получении свойств документа: {e}")

        return properties

    @staticmethod
    def _estimate_page_count(doc) -> int:
        """Приблизительный подсчет страниц"""
        try:
            total_chars = sum(len(para.text) for para in doc.paragraphs)
            estimated_pages = max(1, total_chars // 1250)
            return estimated_pages
        except:
            return 1

    @staticmethod
    def debug_document_fonts(file_path: str) -> Dict:
        """Расширенная отладочная функция для анализа всех шрифтов в документе"""
        try:
            doc = Document(file_path)
            font_analysis = {
                'styles_fonts': {},
                'runs_fonts': {},
                'paragraph_fonts': {},
                'xml_fonts': {},
                'theme_fonts': DocumentLoader._get_theme_fonts(doc),
                'paragraph_count': len(doc.paragraphs),
                'total_runs': 0
            }

            # Анализ шрифтов в стилях
            for style in doc.styles:
                if hasattr(style, 'font') and style.font and style.font.name:
                    font_analysis['styles_fonts'][style.name] = {
                        'name': style.font.name,
                        'size': style.font.size.pt if style.font.size else None
                    }

            # Анализ шрифтов в параграфах и runs
            font_counter = Counter()
            for para_idx, para in enumerate(doc.paragraphs):
                if not para.text.strip():
                    continue

                font_analysis['total_runs'] += len(para.runs)

                # Анализ XML параграфа
                xml_analysis = DocumentLoader._analyze_paragraph_xml(para)
                if xml_analysis['fonts_found']:
                    font_analysis['xml_fonts'][f'para_{para_idx}'] = xml_analysis

                # Анализ стиля параграфа
                if para.style:
                    style_name = para.style.name
                    if style_name not in font_analysis['paragraph_fonts']:
                        font_analysis['paragraph_fonts'][style_name] = []
                    font_analysis['paragraph_fonts'][style_name].append({
                        'para_idx': para_idx,
                        'text_preview': para.text[:50] + ('...' if len(para.text) > 50 else '')
                    })

                # Анализ runs
                for run_idx, run in enumerate(para.runs):
                    if run.text.strip():
                        run_font = None

                        # Проверяем прямое форматирование
                        if run.font and run.font.name:
                            run_font = run.font.name

                        # Проверяем XML
                        if not run_font:
                            run_font = DocumentLoader._get_font_from_xml(run._element)

                        if run_font:
                            font_counter[run_font] += len(run.text)
                            key = f"para_{para_idx}_run_{run_idx}"
                            font_analysis['runs_fonts'][key] = {
                                'font': run_font,
                                'text_length': len(run.text),
                                'text_preview': run.text[:50] + ('...' if len(run.text) > 50 else ''),
                                'source': 'direct' if (run.font and run.font.name) else 'xml'
                            }

            font_analysis['font_frequency'] = dict(font_counter.most_common())
            return font_analysis

        except Exception as e:
            return {'error': str(e)}